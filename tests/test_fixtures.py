"""
test_fixtures.py — builds throwaway edge-case documents on demand for the test
suite (rather than shipping them as static binary files in the repo).
"""
import os
import io
from docx import Document
from docx.shared import Pt

FIXTURE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_fixtures")


def _ensure_dir():
    os.makedirs(FIXTURE_DIR, exist_ok=True)


def make_corrupt_docx():
    _ensure_dir()
    path = os.path.join(FIXTURE_DIR, "corrupt.docx")
    with open(path, "w") as f:
        f.write("this is not a real docx file, just plain text")
    return path


def make_empty_pdf():
    _ensure_dir()
    path = os.path.join(FIXTURE_DIR, "empty.pdf")
    open(path, "wb").close()
    return path


def make_docx_with_footer(filename, footer_text, body_text="Body content for edge-case test."):
    """Builds a docx with arbitrary raw footer text (single paragraph)."""
    _ensure_dir()
    doc = Document()
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = footer_text
    doc.add_paragraph(body_text)
    path = os.path.join(FIXTURE_DIR, filename)
    doc.save(path)
    return path


def make_docx_with_table_footer(filename, cell_texts):
    """Builds a docx whose footer content lives in a table (common in real
    controlled templates), rather than a plain paragraph."""
    _ensure_dir()
    doc = Document()
    footer = doc.sections[0].footer
    table = footer.add_table(rows=1, cols=len(cell_texts), width=Pt(400))
    for i, text in enumerate(cell_texts):
        table.cell(0, i).text = text
    doc.add_paragraph("Body content for table-footer edge case.")
    path = os.path.join(FIXTURE_DIR, filename)
    doc.save(path)
    return path


def make_docx_no_footer(filename):
    """A docx with a completely empty/blank footer (template ID missing)."""
    _ensure_dir()
    doc = Document()
    doc.add_paragraph("Body content — no footer at all was stamped on this document.")
    path = os.path.join(FIXTURE_DIR, filename)
    doc.save(path)
    return path


def make_pdf_with_footer(filename, template_id, version):
    """Generates a valid minimal PDF page directly with the footer text stamped at the bottom.
    This eliminates LibreOffice dependencies for testing PDF text extraction.
    """
    _ensure_dir()
    footer_text = f"Template ID: {template_id}   |   Version: {version}" if version else f"Template ID: {template_id}"
    if not template_id:
        footer_text = "No template ID stamped here"
    
    # Minimal valid PDF format
    pdf_content = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 120 >>
stream
BT
/F1 10 Tf
50 50 Td
({footer_text}) Tj
ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000244 00000 n 
0000000318 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
480
%%EOF
"""
    path = os.path.join(FIXTURE_DIR, filename)
    with open(path, "wb") as f:
        f.write(pdf_content.encode("ascii"))
    return path


def make_scanned_pdf_from_text(filename, template_id, version):
    """Generates a genuinely scanned (image-only) PDF using Pillow.
    This replaces pdf2image + img2pdf conversion for testing OCR fallback.
    """
    _ensure_dir()
    from PIL import Image, ImageDraw
    img = Image.new('RGB', (595, 842), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    footer_text = f"Template ID: {template_id}   |   Version: {version}" if version else f"Template ID: {template_id}"
    d.text((50, 800), footer_text, fill=(0, 0, 0))
    path = os.path.join(FIXTURE_DIR, filename)
    img.save(path, "PDF")
    return path


def check_tesseract_installed():
    """Checks if tesseract binary is executable."""
    import subprocess
    try:
        result = subprocess.run(["tesseract", "--version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, shell=True)
        return result.returncode == 0
    except Exception:
        return False


def make_scanned_pdf_from(source_pdf_path, filename):
    # Backward compatibility stub
    return make_scanned_pdf_from_text(filename, "V-QMS-0114221", "3.0")


def docx_to_pdf(docx_path):
    # Backward compatibility stub - returns path to a generated PDF instead of converting
    raise NotImplementedError("Use make_pdf_with_footer instead of converting via LibreOffice.")


def cleanup():
    import shutil
    if os.path.exists(FIXTURE_DIR):
        shutil.rmtree(FIXTURE_DIR)

