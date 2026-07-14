"""
extractor.py — pulls (template_id, version) out of a .docx or .pdf footer.

Deterministic, regex-based. No AI involved — the ID is a structured string,
not free text, so a model would add cost/latency/uncertainty for zero gain.
OCR is only reached for scanned PDFs with no extractable text layer.
"""
import re
import os

# Matches: "V-QMS-0114221" ... "Version: 3.0" (or "Ver", "v", "Rev", "V3.0" etc.,
# any separator/punctuation, any order of words in between). Case-insensitive
# since footers are sometimes typed in lowercase or ALL CAPS.
ID_PATTERN = re.compile(r"(V-QMS-\d{7})", re.IGNORECASE)
VERSION_PATTERN = re.compile(r"(?:Version|Revision|Ver|Rev|V)\.?\s{0,2}:?\s{0,2}(\d+(?:\.\d+)?)", re.IGNORECASE)


def _parse(text: str):
    id_match = ID_PATTERN.search(text)
    ver_match = VERSION_PATTERN.search(text)
    # Normalize to uppercase — master list IDs are canonically uppercase, and a
    # footer typed/OCR'd in lowercase shouldn't cause a false "not found".
    template_id = id_match.group(1).upper() if id_match else None
    version = ver_match.group(1) if ver_match else None
    return template_id, version


def extract_from_docx(path: str):
    """Reads all section footers of a .docx and returns (template_id, version, raw_text)."""
    from docx import Document
    doc = Document(path)
    footer_text = ""
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is None:
                continue
            for p in footer.paragraphs:
                footer_text += p.text + "\n"
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        footer_text += cell.text + "\n"
    template_id, version = _parse(footer_text)
    return template_id, version, footer_text.strip()


def extract_from_pdf(path: str):
    """Reads the bottom strip of each PDF page. Falls back to OCR if no text layer found."""
    import pdfplumber
    footer_text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            h = page.height
            # bottom 12% of the page — where footers live
            crop = page.within_bbox((0, h * 0.88, page.width, h))
            footer_text += (crop.extract_text() or "") + "\n"

    template_id, version = _parse(footer_text)
    if template_id:
        return template_id, version, footer_text.strip()

    # No usable text layer -> likely a scanned PDF. OCR fallback.
    return _ocr_fallback(path)


def _ocr_fallback(path: str):
    """OCR fallback for scanned/flattened PDFs with no text layer."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        return None, None, "[OCR unavailable — install pytesseract + pdf2image + poppler]"

    pages = convert_from_path(path, dpi=200)
    footer_text = ""
    for page_img in pages:
        w, h = page_img.size
        crop = page_img.crop((0, int(h * 0.88), w, h))
        footer_text += pytesseract.image_to_string(crop) + "\n"

    template_id, version = _parse(footer_text)
    return template_id, version, footer_text.strip()


def extract(path: str):
    """Entry point — dispatches by file extension. Returns dict result.
    Never raises: corrupt/unreadable/empty files come back as a graceful error
    dict so the caller (UI or batch job) can flag it instead of crashing.
    """
    if not os.path.exists(path):
        return {"template_id": None, "version": None, "raw_footer": None,
                "error": f"File not found: {path}"}

    ext = os.path.splitext(path)[1].lower()
    if ext not in (".docx", ".pdf"):
        return {"template_id": None, "version": None, "raw_footer": None,
                "error": f"Unsupported file type: {ext}"}

    try:
        if ext == ".docx":
            template_id, version, raw = extract_from_docx(path)
        else:
            template_id, version, raw = extract_from_pdf(path)
    except Exception as e:
        return {"template_id": None, "version": None, "raw_footer": None,
                "error": f"Could not read file — it may be corrupt, password-protected, "
                         f"or not a valid {ext} file. ({type(e).__name__}: {e})"}

    return {"template_id": template_id, "version": version, "raw_footer": raw, "error": None}


if __name__ == "__main__":
    import sys
    result = extract(sys.argv[1])
    print(result)
