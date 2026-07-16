"""
Builds 4 sample .docx files mimicking real controlled-template output,
each with a Template ID + Version stamped in the footer — covering all
four verdicts the checker can produce.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(SCRIPT_DIR, "sample_docs")
os.makedirs(OUT_DIR, exist_ok=True)


def add_footer(section, template_id, version):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run(f"Template ID: {template_id}   |   Version: {version}   |   CONFIDENTIAL - Alcon Internal Use Only")
    run.font.size = Pt(8)
    run.font.name = "Arial"


def build_doc(filename, title, template_id, version, body_note):
    doc = Document()
    section = doc.sections[0]
    add_footer(section, template_id, version)

    h = doc.add_heading(title, level=1)
    doc.add_paragraph(
        "Document Number: DOC-2026-0042        Effective Date: 14-Jul-2026        "
        "Prepared by: K. Choudhary, Software Engineer Intern"
    )
    doc.add_paragraph(
        "1. Purpose\nThis document describes the deviation investigation process for "
        "manufacturing non-conformances identified during in-process inspection."
    )
    doc.add_paragraph(
        "2. Scope\nApplies to all manufacturing sites producing Class II/III devices "
        "under the Alcon Quality Management System."
    )
    doc.add_paragraph(f"3. Note\n{body_note}")
    doc.save(os.path.join(OUT_DIR, filename))
    print(f"Wrote {filename}  (footer: {template_id}, v{version})")


# 1. PASS — Effective template, latest version
build_doc(
    "SOP_Deviation_Latest.docx",
    "Deviation Management SOP",
    "V-QMS-0114221", "3.0",
    "This document is authored on the current approved template version.",
)

# 2. WARN — Effective template, but an older version than latest (3.0)
build_doc(
    "SOP_Deviation_Outdated_Version.docx",
    "Deviation Management SOP",
    "V-QMS-0114221", "2.0",
    "This document was drafted some months ago and may be on an older template version.",
)

# 3. FAIL — fully superseded Template ID (V-QMS-0114100 -> replaced by V-QMS-0114221)
build_doc(
    "SOP_Deviation_Superseded_Template.docx",
    "Deviation Management SOP",
    "V-QMS-0114100", "2.0",
    "This document appears to use a template carried over from a shared drive archive.",
)

# 4. FLAG — Template ID not present in master list at all
build_doc(
    "Form_Unknown_Template.docx",
    "Unknown Internal Form",
    "V-QMS-9999999", "1.0",
    "This document uses a template ID that does not appear in the master list.",
)
